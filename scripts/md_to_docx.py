#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
from pathlib import Path
import re

md = Path('docs/PMOMax_Business_Pricing.md')
if not md.exists():
    print('Source markdown not found:', md)
    raise SystemExit(1)

text = md.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

list_mode = False
for line in lines:
    if line.strip() == '':
        # blank line
        list_mode = False
        doc.add_paragraph('')
        continue
    # Headings
    m = re.match(r'^(#{1,6})\s+(.*)$', line)
    if m:
        level = len(m.group(1))
        heading = m.group(2).strip()
        # map: 1->0 (Title), 2->1, else -> 2
        if level == 1:
            doc.add_heading(heading, level=0)
        elif level == 2:
            doc.add_heading(heading, level=1)
        else:
            doc.add_heading(heading, level=2)
        list_mode = False
        continue
    # Unordered list
    m = re.match(r'^\s*[-*+]\s+(.*)$', line)
    if m:
        text_item = m.group(1).strip()
        p = doc.add_paragraph(text_item)
        p.style = 'List Bullet'
        list_mode = True
        continue
    # Numbered list
    m = re.match(r'^\s*\d+[.)]\s+(.*)$', line)
    if m:
        text_item = m.group(1).strip()
        p = doc.add_paragraph(text_item)
        p.style = 'List Number'
        list_mode = True
        continue
    # Blockquote
    m = re.match(r'^>\s?(.*)$', line)
    if m:
        p = doc.add_paragraph(m.group(1).strip())
        p.style = 'Intense Quote'
        continue
    # Horizontal rule
    if re.match(r'^---+$', line):
        doc.add_page_break()
        continue
    # Default paragraph
    if list_mode:
        # continue list if previous was list—just add another bullet
        p = doc.add_paragraph(line.strip())
        p.style = 'List Bullet'
    else:
        doc.add_paragraph(line.strip())

out = Path('docs/PMOMax_Business_Pricing.docx')
doc.save(out)
print('Wrote', out)
